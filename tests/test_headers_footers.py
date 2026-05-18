"""Tests for header/footer extraction and roundtrip."""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from arbordoc.core.extractor import extract_headers_footers, extract_all, extract_section_properties
from arbordoc.core.parser import parse_docx
from arbordoc.core.styler import transform_docx
from arbordoc.models.schema import BlockType, NodeType


def _add_header_text(section, text: str) -> None:
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = text if header.paragraphs else header.add_paragraph(text)


def _add_footer_text(section, text: str) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = text if footer.paragraphs else footer.add_paragraph(text)


def test_extract_header(tmp_path) -> None:
    doc_path = tmp_path / "header.docx"
    doc = Document()
    _add_header_text(doc.sections[0], "Company Confidential")
    doc.add_paragraph("Body text")
    doc.save(doc_path)

    loaded = Document(doc_path)
    hf_blocks = extract_headers_footers(loaded)

    headers = [b for b in hf_blocks if b.type == BlockType.HEADER]
    assert len(headers) >= 1
    assert any("Company Confidential" in (b.text or "") for b in headers)


def test_extract_footer(tmp_path) -> None:
    doc_path = tmp_path / "footer.docx"
    doc = Document()
    _add_footer_text(doc.sections[0], "Page 1 of 10")
    doc.add_paragraph("Body text")
    doc.save(doc_path)

    loaded = Document(doc_path)
    hf_blocks = extract_headers_footers(loaded)

    footers = [b for b in hf_blocks if b.type == BlockType.FOOTER]
    assert len(footers) >= 1


def test_extract_section_properties(tmp_path) -> None:
    doc_path = tmp_path / "sections.docx"
    doc = Document()
    doc.save(doc_path)

    loaded = Document(doc_path)
    sections = extract_section_properties(loaded)
    assert len(sections) >= 1
    assert "page_width" in sections[0]
    assert "page_height" in sections[0]


def test_header_roundtrip(sample_template_docx, tmp_path) -> None:
    src_path = tmp_path / "header_src.docx"
    doc = Document()
    _add_header_text(doc.sections[0], "HEADER TEXT")
    doc.add_paragraph("Body content")
    doc.save(src_path)

    out_path = tmp_path / "header_out.docx"
    transform_docx(src_path, sample_template_docx, out_path)

    rendered = Document(out_path)
    header = rendered.sections[0].header
    header_text = " ".join(p.text for p in header.paragraphs)
    assert "HEADER TEXT" in header_text


def test_footer_roundtrip(sample_template_docx, tmp_path) -> None:
    src_path = tmp_path / "footer_src.docx"
    doc = Document()
    _add_footer_text(doc.sections[0], "FOOTER TEXT")
    doc.add_paragraph("Body content")
    doc.save(src_path)

    out_path = tmp_path / "footer_out.docx"
    transform_docx(src_path, sample_template_docx, out_path)

    rendered = Document(out_path)
    footer = rendered.sections[0].footer
    footer_text = " ".join(p.text for p in footer.paragraphs)
    assert "FOOTER TEXT" in footer_text


def test_doc_without_header_footer_works_fine(sample_source_docx, sample_template_docx, tmp_path) -> None:
    """Documents without headers/footers should still render correctly."""
    out_path = tmp_path / "no_hf.docx"
    transform_docx(sample_source_docx, sample_template_docx, out_path)
    assert out_path.is_file()
    rendered = Document(out_path)
    assert len(rendered.paragraphs) >= 1
