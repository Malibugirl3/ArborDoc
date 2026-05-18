from __future__ import annotations

import base64
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

# Valid 1×1 PNG for embedding in test documents.
_MINIMAL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQYBACHV7NkAAAAASUVORK5CYII="
)


@pytest.fixture()
def sample_docx_with_image(tmp_path: Path) -> Path:
    """DOCX with one inline picture (paragraph + image)."""
    png_path = tmp_path / "tiny.png"
    png_path.write_bytes(_MINIMAL_PNG)
    docx_path = tmp_path / "with_image.docx"
    document = Document()
    document.add_paragraph("Has image:")
    document.add_picture(str(png_path), width=Inches(0.5))
    document.save(docx_path)
    return docx_path


@pytest.fixture()
def sample_source_docx(tmp_path: Path) -> Path:
    path = tmp_path / "source.docx"
    document = Document()
    document.add_heading("Main Title", level=1)
    document.add_paragraph("Intro paragraph.")
    document.add_heading("Section A", level=2)
    document.add_paragraph("Section body.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    document.save(path)
    return path


@pytest.fixture()
def sample_template_docx(tmp_path: Path) -> Path:
    path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Template placeholder.")
    document.save(path)
    return path


@pytest.fixture()
def sample_formatted_docx(tmp_path: Path) -> Path:
    """DOCX with bold, italic, colored runs and basic paragraph formatting."""
    from docx.shared import Pt, RGBColor

    path = tmp_path / "formatted.docx"
    document = Document()

    para = document.add_paragraph()
    run = para.add_run("Bold text")
    run.font.bold = True
    run = para.add_run(" normal ")
    run = para.add_run("Italic text")
    run.font.italic = True

    para2 = document.add_paragraph()
    run = para2.add_run("Blue text")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    run = para2.add_run(" and ")
    run = para2.add_run("Red text")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    para3 = document.add_paragraph()
    para3.paragraph_format.alignment = 1  # CENTER
    para3.add_run("Centered paragraph")

    document.save(path)
    return path
