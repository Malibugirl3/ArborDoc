"""Tests for numbered/bulleted list extraction and roundtrip."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from arbordoc.core.extractor import extract_blocks, extract_list_info
from arbordoc.core.parser import parse_docx
from arbordoc.core.styler import transform_docx


def _make_list_docx(path: Path, ordered: bool) -> Path:
    """Create a DOCX with a real numbered/bulleted list using XML."""
    document = Document()
    body = document.element.body

    # Add a heading
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "Heading1")
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "List Test"
    r.append(t)
    p.append(r)
    body.append(p)

    num_id = 42  # avoid conflicts with built-in numbering (1-9)
    abstract_num_id = 42

    # Create numbering part
    numbering_part = document.part.numbering_part
    numbering_el = numbering_part._element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))
    lvl0 = OxmlElement("w:lvl")
    lvl0.set(qn("w:ilvl"), "0")
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl0.append(numFmt)
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "%1." if ordered else "•")
    lvl0.append(lvlText)
    abstract_num.append(lvl0)
    numbering_el.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abs_ref)
    numbering_el.append(num)

    # Add list items
    for i in range(3):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), "0")
        numPr.append(ilvl_el)
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(num_id))
        numPr.append(numId_el)
        pPr.append(numPr)
        p.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = f"Item {i + 1}"
        r.append(t)
        p.append(r)
        body.append(p)

    document.save(path)
    return path


def test_extract_bullet_list(tmp_path) -> None:
    path = _make_list_docx(tmp_path / "bullet.docx", ordered=False)
    doc = Document(path)

    # Extract and find list paragraphs
    li = extract_list_info(doc.paragraphs[1])  # First list item
    assert li is not None
    assert li["num_id"] == 42
    assert li["level"] == 0


def test_extract_numbered_list(tmp_path) -> None:
    path = _make_list_docx(tmp_path / "numbered.docx", ordered=True)
    doc = Document(path)

    li = extract_list_info(doc.paragraphs[1])
    assert li is not None
    assert li["num_id"] == 42


def test_list_roundtrip_bullet(sample_template_docx, tmp_path) -> None:
    source_path = _make_list_docx(tmp_path / "bullet_src.docx", ordered=False)
    out_path = tmp_path / "bullet_out.docx"

    transform_docx(source_path, sample_template_docx, out_path)

    rendered = Document(out_path)
    paragraphs = rendered.paragraphs
    # Should have heading + 3 list items
    list_texts = [p.text for p in paragraphs if p.text.startswith("Item")]
    assert len(list_texts) == 3
    assert list_texts[0] == "Item 1"
    assert list_texts[1] == "Item 2"
    assert list_texts[2] == "Item 3"


def test_list_roundtrip_numbered(sample_template_docx, tmp_path) -> None:
    source_path = _make_list_docx(tmp_path / "numbered_src.docx", ordered=True)
    out_path = tmp_path / "numbered_out.docx"

    transform_docx(source_path, sample_template_docx, out_path)

    rendered = Document(out_path)
    list_texts = [p.text for p in rendered.paragraphs if p.text.startswith("Item")]
    assert len(list_texts) == 3


def test_list_info_in_parsed_tree(tmp_path) -> None:
    path = _make_list_docx(tmp_path / "tree_test.docx", ordered=True)
    root = parse_docx(path)

    list_nodes = [
        n for n in root.iter_depth_first()
        if n.list_info is not None
    ]
    assert len(list_nodes) == 3
    assert list_nodes[0].list_info.is_ordered is True
    assert list_nodes[0].list_info.level == 0


def test_non_list_paragraph_has_no_list_info(sample_source_docx) -> None:
    root = parse_docx(sample_source_docx)
    for node in root.iter_depth_first():
        assert node.list_info is None
