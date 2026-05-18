"""Tests for run formatting, hyperlinks, and inline content roundtrip."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor

from arbordoc.core.extractor import (
    extract_blocks,
    extract_paragraph_format,
    extract_run_format,
    extract_runs_from_paragraph,
)
from arbordoc.core.parser import build_tree_from_blocks, parse_docx
from arbordoc.core.styler import render_tree_to_template, transform_docx
from arbordoc.converters.latex import LatexExporter
from arbordoc.models.schema import (
    DocNode,
    HyperlinkRun,
    InlineElementType,
    InlineImageInline,
    NodeType,
    RunFormat,
    TextRun,
)


def test_extract_run_bold_italic(sample_formatted_docx) -> None:
    document = Document(sample_formatted_docx)
    blocks = extract_blocks(document)

    runs = blocks[0].meta["runs"]
    assert len(runs) == 3
    assert runs[0]["text"] == "Bold text"
    assert runs[0]["format"]["bold"] is True
    assert runs[1]["text"] == " normal "
    assert runs[2]["text"] == "Italic text"
    assert runs[2]["format"]["italic"] is True


def test_extract_run_font_color(sample_formatted_docx) -> None:
    document = Document(sample_formatted_docx)
    blocks = extract_blocks(document)

    runs = blocks[1].meta["runs"]
    assert len(runs) == 3
    assert runs[0]["format"]["font_color_rgb"] == "0000FF"
    assert runs[2]["format"]["font_color_rgb"] == "FF0000"


def test_extract_paragraph_format_alignment(sample_formatted_docx) -> None:
    document = Document(sample_formatted_docx)
    blocks = extract_blocks(document)

    pf = blocks[2].meta["paragraph_format"]
    assert pf["alignment"] == "CENTER"


def test_parse_builds_inline_content(sample_formatted_docx) -> None:
    root = parse_docx(sample_formatted_docx)

    para_nodes = [n for n in root.children if n.type == NodeType.PARAGRAPH]
    assert len(para_nodes) >= 3

    first = para_nodes[0]
    assert first.inline_content is not None
    assert len(first.inline_content) >= 3

    text_runs = [e for e in first.inline_content if isinstance(e, TextRun)]
    assert text_runs[0].text == "Bold text"
    assert text_runs[0].format.bold is True
    assert text_runs[2].text == "Italic text"
    assert text_runs[2].format.italic is True


def test_render_formatting_roundtrip(sample_formatted_docx, sample_template_docx, tmp_path) -> None:
    output_path = tmp_path / "formatted_roundtrip.docx"
    transform_docx(sample_formatted_docx, sample_template_docx, output_path)

    rendered = Document(output_path)
    paragraphs = rendered.paragraphs

    assert len(paragraphs) >= 3

    first_para_runs = paragraphs[0].runs
    assert len(first_para_runs) == 3
    assert first_para_runs[0].text == "Bold text"
    assert first_para_runs[0].font.bold is True
    assert first_para_runs[2].text == "Italic text"
    assert first_para_runs[2].font.italic is True

    third_para = paragraphs[2]
    assert third_para.paragraph_format.alignment == 1  # CENTER


def test_render_fallback_to_plain_text(sample_source_docx, sample_template_docx, tmp_path) -> None:
    """DocNode without inline_content still renders via text field (backward compat)."""
    output_path = tmp_path / "fallback.docx"
    transform_docx(sample_source_docx, sample_template_docx, output_path)

    rendered = Document(output_path)
    texts = [p.text for p in rendered.paragraphs]
    assert "Main Title" in texts
    assert "Intro paragraph." in texts


def test_docnode_with_inline_content_renders_formatting(
    sample_template_docx, tmp_path
) -> None:
    """Manually constructed DocNode with inline_content renders formatted runs."""
    root = DocNode(
        type=NodeType.DOCUMENT,
        children=[
            DocNode(
                type=NodeType.PARAGRAPH,
                level=1,
                text="Bold and Italic",
                inline_content=[
                    TextRun(text="Bold ", format=RunFormat(bold=True)),
                    TextRun(text="Italic", format=RunFormat(italic=True)),
                ],
            ),
        ],
    )
    output_path = tmp_path / "manual_inline.docx"

    render_tree_to_template(root, sample_template_docx, output_path)

    rendered = Document(output_path)
    runs = rendered.paragraphs[0].runs
    assert len(runs) == 2
    assert runs[0].text == "Bold "
    assert runs[0].font.bold is True
    assert runs[1].text == "Italic"
    assert runs[1].font.italic is True


def test_latex_inline_formatting_bold_italic() -> None:
    root = DocNode(
        type=NodeType.DOCUMENT,
        children=[
            DocNode(
                type=NodeType.PARAGRAPH,
                level=1,
                text="BoldItalic",
                inline_content=[
                    TextRun(text="Bold", format=RunFormat(bold=True)),
                    TextRun(text=" and "),
                    TextRun(text="Italic", format=RunFormat(italic=True)),
                ],
            ),
        ],
    )
    rendered = LatexExporter(standalone=False).export(root)
    assert r"\textbf{Bold}" in rendered
    assert r"\textit{Italic}" in rendered


def test_latex_fallback_to_plain_text() -> None:
    """When inline_content is None, LaTeX exporter uses plain text."""
    root = DocNode(
        type=NodeType.DOCUMENT,
        children=[
            DocNode(type=NodeType.PARAGRAPH, level=1, text="Plain text"),
        ],
    )
    rendered = LatexExporter(standalone=False).export(root)
    assert "Plain text" in rendered


def test_hyperlink_in_inline_content() -> None:
    root = DocNode(
        type=NodeType.DOCUMENT,
        children=[
            DocNode(
                type=NodeType.PARAGRAPH,
                level=1,
                text="Click here",
                inline_content=[
                    TextRun(text="Before "),
                    HyperlinkRun(
                        url="https://example.com",
                        children=[TextRun(text="Click here")],
                    ),
                    TextRun(text=" after"),
                ],
            ),
        ],
    )
    rendered = LatexExporter(standalone=False).export(root)
    assert r"\href{https://example.com}{Click here}" in rendered


def test_inline_image_in_content() -> None:
    root = DocNode(
        type=NodeType.DOCUMENT,
        children=[
            DocNode(
                type=NodeType.PARAGRAPH,
                level=1,
                text="Image:",
                inline_content=[
                    InlineImageInline(relationship_id="rId42"),
                ],
            ),
        ],
    )
    rendered = LatexExporter(standalone=False).export(root)
    assert "rId42" in rendered
    assert "% TODO: inline image" in rendered


def test_run_extraction_strikethrough(tmp_path) -> None:
    doc_path = tmp_path / "strike.docx"
    document = Document()
    para = document.add_paragraph()
    run = para.add_run("Struck")
    run.font.strike = True
    document.save(doc_path)

    doc = Document(doc_path)
    blocks = extract_blocks(doc)
    runs = blocks[0].meta["runs"]
    assert runs[0]["format"]["strikethrough"] is True


def test_run_extraction_superscript_subscript(tmp_path) -> None:
    doc_path = tmp_path / "super_sub.docx"
    document = Document()
    para = document.add_paragraph()
    run = para.add_run("Super")
    run.font.superscript = True
    run2 = para.add_run("Sub")
    run2.font.subscript = True
    document.save(doc_path)

    doc = Document(doc_path)
    blocks = extract_blocks(doc)
    runs = blocks[0].meta["runs"]
    assert runs[0]["format"]["superscript"] is True
    assert runs[1]["format"]["subscript"] is True


def test_run_extraction_font_name_size(tmp_path) -> None:
    doc_path = tmp_path / "font.docx"
    document = Document()
    para = document.add_paragraph()
    run = para.add_run("Big Arial")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    document.save(doc_path)

    doc = Document(doc_path)
    blocks = extract_blocks(doc)
    runs = blocks[0].meta["runs"]
    assert runs[0]["format"]["font_name"] == "Arial"
    assert runs[0]["format"]["font_size_pt"] == 18.0
