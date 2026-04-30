from docx import Document

from arbordoc.core.styler import render_tree_to_template, transform_docx
from arbordoc.models.schema import DocNode, NodeType


def test_render_tree_to_template_writes_supported_content(sample_template_docx, tmp_path) -> None:
    root = DocNode(
        type=NodeType.DOCUMENT,
        children=[
            DocNode(type=NodeType.HEADING, level=1, text="Rendered Title"),
            DocNode(type=NodeType.PARAGRAPH, level=1, text="Rendered paragraph"),
            DocNode(type=NodeType.TABLE, level=1, meta={"rows": [["ignored"]]}),
        ],
    )
    output_path = tmp_path / "rendered.docx"

    render_tree_to_template(root, sample_template_docx, output_path)

    document = Document(output_path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert texts == ["Rendered Title", "Rendered paragraph"]
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "ignored"


def test_transform_docx_creates_output(sample_source_docx, sample_template_docx, tmp_path) -> None:
    output_path = tmp_path / "transformed.docx"

    transform_docx(sample_source_docx, sample_template_docx, output_path)

    document = Document(output_path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Main Title" in texts
    assert "Intro paragraph." in texts
    assert len(document.tables) >= 1
    assert document.tables[0].cell(0, 0).text == "A1"


def test_transform_embedded_image_roundtrip(sample_docx_with_image, sample_template_docx, tmp_path) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    output_path = tmp_path / "with_pic_out.docx"

    transform_docx(sample_docx_with_image, sample_template_docx, output_path)

    rendered = Document(output_path)
    image_rels = sum(1 for r in rendered.part.rels.values() if r.reltype == RT.IMAGE)
    assert image_rels >= 1
