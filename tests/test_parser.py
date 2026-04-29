from docx import Document

from arbordoc.core.extractor import extract_blocks
from arbordoc.core.parser import build_tree_from_blocks, parse_document, parse_docx
from arbordoc.core.tree import walk_depth_first
from arbordoc.models.schema import BlockType, NodeType


def test_extractor_normalizes_docx_into_blocks(sample_source_docx) -> None:
    document = Document(sample_source_docx)

    blocks = extract_blocks(document)

    assert [block.type for block in blocks] == [
        BlockType.HEADING,
        BlockType.PARAGRAPH,
        BlockType.HEADING,
        BlockType.PARAGRAPH,
        BlockType.TABLE,
    ]
    assert blocks[0].level == 1
    assert blocks[2].level == 2


def test_parser_builds_heading_stack(sample_source_docx) -> None:
    document = Document(sample_source_docx)
    root = build_tree_from_blocks(extract_blocks(document), source_path=str(sample_source_docx))

    assert len(root.children) == 1

    top_heading = root.children[0]
    assert top_heading.type == NodeType.HEADING
    assert top_heading.text == "Main Title"

    nested_heading = next(node for node in top_heading.children if node.type == NodeType.HEADING)
    assert nested_heading.level == 2
    assert nested_heading.text == "Section A"


def test_parser_preserves_paragraph_and_table_order() -> None:
    document = Document()
    document.add_heading("Title", level=1)
    document.add_paragraph("First paragraph")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    document.add_paragraph("Second paragraph")

    root = parse_document(document)
    heading = root.children[0]
    content_types = [child.type for child in heading.children]

    assert content_types == [NodeType.PARAGRAPH, NodeType.TABLE, NodeType.PARAGRAPH]


def test_walk_depth_first_returns_expected_nodes(sample_source_docx) -> None:
    root = parse_docx(sample_source_docx)
    node_types = [node.type for node in walk_depth_first(root)]

    assert node_types[0] == NodeType.DOCUMENT
    assert NodeType.HEADING in node_types
    assert NodeType.TABLE in node_types
