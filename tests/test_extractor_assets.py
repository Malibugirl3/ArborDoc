"""Tests for image blob helpers in ``arbordoc.core.extractor``."""

from __future__ import annotations

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from arbordoc.core.extractor import (
    extract_image_blob_cache,
    extract_image_blob_cache_to_directory,
    get_image_blob,
)


def test_extract_image_blob_cache_roundtrip(sample_docx_with_image) -> None:
    document = Document(sample_docx_with_image)
    cache = extract_image_blob_cache(document)

    assert len(cache) >= 1
    rid, blob = next(iter(cache.items()))
    assert isinstance(blob, bytes)
    assert len(blob) >= 8
    assert get_image_blob(document, rid) == blob


def test_extract_image_blob_cache_to_directory_writes_files(sample_docx_with_image, tmp_path) -> None:
    document = Document(sample_docx_with_image)
    out_dir = tmp_path / "img_dump"
    base, path_map = extract_image_blob_cache_to_directory(document, out_dir)

    assert base.resolve() == out_dir.resolve()
    assert path_map
    for path in path_map.values():
        assert path.is_file()
        assert path.stat().st_size > 0


def test_png_document_contains_image_relationship(sample_docx_with_image) -> None:
    doc = Document(sample_docx_with_image)
    reltypes = [r.reltype for r in doc.part.rels.values()]

    assert RT.IMAGE in reltypes
