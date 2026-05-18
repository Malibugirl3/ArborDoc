"""Tests for FastAPI endpoints."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from arbordoc.api.app import app

client = TestClient(app)


@pytest.fixture()
def source_docx(tmp_path: Path) -> Path:
    path = tmp_path / "source.docx"
    doc = Document()
    doc.add_heading("Test API", level=1)
    doc.add_paragraph("Body content.")
    doc.save(str(path))
    return path


@pytest.fixture()
def template_docx(tmp_path: Path) -> Path:
    path = tmp_path / "template.docx"
    doc = Document()
    doc.save(str(path))
    return path


def test_health():
    resp = client.get("/health")
    assert resp.status_code == 200
    assert resp.json() == {"status": "ok"}


def test_parse_returns_json(source_docx: Path):
    with open(source_docx, "rb") as f:
        resp = client.post("/parse", files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert resp.status_code == 200
    data = resp.json()
    assert data["type"] == "document"
    assert len(data["children"]) >= 1


def test_parse_rejects_non_docx():
    resp = client.post("/parse", files={"file": ("test.txt", b"not docx", "text/plain")})
    assert resp.status_code == 400
    assert "error" in resp.json()


def test_transform_returns_docx(source_docx: Path, template_docx: Path, tmp_path: Path):
    with open(source_docx, "rb") as sf, open(template_docx, "rb") as tf:
        resp = client.post(
            "/transform",
            files={
                "source": ("source.docx", sf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                "template": ("template.docx", tf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            },
        )
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    # Verify it's a valid DOCX by reading it
    out = tmp_path / "api_output.docx"
    out.write_bytes(resp.content)
    doc = Document(str(out))
    assert len(doc.paragraphs) >= 1


def test_transform_rejects_non_docx():
    resp = client.post(
        "/transform",
        files={
            "source": ("source.txt", b"not docx", "text/plain"),
            "template": ("template.txt", b"not docx", "text/plain"),
        },
    )
    assert resp.status_code == 400


def test_doc_with_header_footer_roundtrip(tmp_path: Path):
    """Full roundtrip via API: source with header/footer through transform."""
    src = tmp_path / "hf_source.docx"
    doc = Document()
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "API Header"
    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = "API Footer"
    doc.add_paragraph("Body")
    doc.save(str(src))

    tpl = tmp_path / "hf_template.docx"
    Document().save(str(tpl))

    with open(src, "rb") as sf, open(tpl, "rb") as tf:
        resp = client.post(
            "/transform",
            files={
                "source": ("source.docx", sf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                "template": ("template.docx", tf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            },
        )

    assert resp.status_code == 200
    out = tmp_path / "hf_output.docx"
    out.write_bytes(resp.content)
    rendered = Document(str(out))
    hdr_text = " ".join(p.text for p in rendered.sections[0].header.paragraphs)
    assert "API Header" in hdr_text
