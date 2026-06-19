"""
@file styler.py
@brief Rebuild DOCX documents from DocTree using a template.

@author Ma PingChuan, Shi Kaibo
@copyright Copyright (c) 2026 Ma PingChuan, Shi Kaibo. SPDX-License-Identifier: MIT
@date 2026

Takes a parsed DocTree and writes supported content back into a template document.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Dict, List, Mapping, Optional, Sequence, Union, cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from arbordoc.core.extractor import extract_image_blob_cache
from arbordoc.core.parser import parse_docx
from arbordoc.core.tree import walk_depth_first
from arbordoc.models.schema import (
    DocNode,
    HyperlinkRun,
    InlineImageInline,
    InlineElementType,
    NodeType,
    TextRun,
)

DEFAULT_STYLE_MAP = {
    "paragraph": "Normal",
}

DEFAULT_IMAGE_RENDER_WIDTH = Inches(4)


def _table_rows_as_matrix(rows: object) -> List[List[str]]:
    """Flatten cell values from parser meta into rectangular text rows."""
    if not isinstance(rows, Sequence) or isinstance(rows, (str, bytes)):
        return []

    normalized: List[List[str]] = []
    for row in rows:
        if isinstance(row, Sequence) and not isinstance(row, (str, bytes)):
            normalized.append([str(cell) if cell is not None else "" for cell in cast(Sequence[object], row)])
        else:
            normalized.append(["" if row is None else str(row)])

    width = max((len(row) for row in normalized), default=0)
    for row in normalized:
        if len(row) < width:
            row.extend([""] * (width - len(row)))
    return normalized


def _append_table(document: DocumentObject, meta: Mapping[str, object]) -> None:
    cells = meta.get("cells")
    rows = _table_rows_as_matrix(meta.get("rows"))
    if not rows:
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = ""
        return

    nrows = len(rows)
    ncols = len(rows[0])
    table = document.add_table(rows=nrows, cols=ncols)

    # Apply column widths if available
    column_widths = meta.get("column_widths")
    if column_widths and isinstance(column_widths, Sequence) and not isinstance(column_widths, (str, bytes)):
        for j, width in enumerate(column_widths):
            if j < ncols:
                try:
                    table.columns[j].width = int(float(str(width)))
                except (ValueError, TypeError):
                    pass

    # Apply table style
    table_style = meta.get("table_style")
    if table_style and isinstance(table_style, str):
        try:
            table.style = table_style
        except (KeyError, ValueError):
            pass

    # Process cells (handling grid_span and v_merge)
    v_merge_restarts: dict[int, int] = {}
    for i, row in enumerate(rows):
        col_j = 0
        row_cells = cells[i] if cells and i < len(cells) and isinstance(cells[i], Sequence) else None
        for j in range(len(row)):
            cell = table.cell(i, col_j)

            # Set text
            if j < len(row):
                cell.text = row[j]

            # Handle grid_span
            gs = 1
            if row_cells and j < len(row_cells):
                cell_meta = row_cells[j]
                if isinstance(cell_meta, dict):
                    gs_raw = cell_meta.get("grid_span")
                    if isinstance(gs_raw, int) and gs_raw > 1:
                        gs = gs_raw
                    v_merge_val = cell_meta.get("v_merge")
                    if v_merge_val == "restart":
                        try:
                            v_merge_restarts[col_j] = i
                            cell.merge(table.cell(i + 1, col_j))
                        except Exception:
                            pass
                    elif v_merge_val == "continue":
                        pass  # already handled by restart's merge

            if gs > 1 and col_j + gs <= ncols:
                try:
                    cell.merge(table.cell(i, col_j + gs - 1))
                except Exception:
                    pass

            col_j += gs
            if col_j >= ncols:
                break


def _append_image(
    document: DocumentObject,
    blob: bytes,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(
        BytesIO(blob),
        width=DEFAULT_IMAGE_RENDER_WIDTH,
    )


def _clear_document_body(document: DocumentObject) -> None:
    """Remove existing body content while keeping section settings intact."""
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _safe_apply_style(paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except (KeyError, ValueError):
        return


def _heading_style_name(level: int, style_map: Dict[str, str]) -> str:
    return style_map.get(f"heading.{level}") or style_map.get("heading") or f"Heading {level}"


def _apply_run_format(run, fmt) -> None:
    """Apply RunFormat fields to a python-docx Run."""
    if not fmt:
        return
    for attr, value in fmt.items():
        if value is None:
            continue
        if attr == "bold":
            run.font.bold = value
        elif attr == "italic":
            run.font.italic = value
        elif attr == "underline":
            run.font.underline = value
        elif attr == "strikethrough":
            run.font.strike = value
        elif attr == "superscript":
            run.font.superscript = value
        elif attr == "subscript":
            run.font.subscript = value
        elif attr == "font_name":
            run.font.name = value
        elif attr == "font_size_pt":
            run.font.size = Pt(float(value))
        elif attr == "font_color_rgb":
            try:
                run.font.color.rgb = RGBColor.from_string(str(value))
            except (ValueError, TypeError):
                pass
        elif attr == "highlight_color":
            run.font.highlight_color = value


def _add_paragraph_with_inline_content(
    document: DocumentObject,
    node: DocNode,
    image_blob_cache: Dict[str, bytes],
) -> object:
    """Create a paragraph from inline_content with full formatting."""
    paragraph = document.add_paragraph()
    inline = node.inline_content or []

    for element in inline:
        if isinstance(element, TextRun):
            run = paragraph.add_run(element.text)
            fmt = element.format.model_dump(exclude_none=True) if element.format else {}
            _apply_run_format(run, fmt)
        elif isinstance(element, HyperlinkRun):
            for child_run_data in element.children:
                run = paragraph.add_run(child_run_data.text)
                child_fmt = child_run_data.format.model_dump(exclude_none=True) if child_run_data.format else {}
                _apply_run_format(run, child_fmt)
        elif isinstance(element, InlineImageInline):
            blob = image_blob_cache.get(element.relationship_id)
            if blob is not None:
                paragraph.add_run().add_picture(BytesIO(blob), width=DEFAULT_IMAGE_RENDER_WIDTH)

    return paragraph


def _apply_paragraph_formatting(paragraph, node: DocNode) -> None:
    """Apply paragraph-level formatting from node.paragraph_format or meta fallback."""
    node_pf = node.paragraph_format
    if node_pf is not None:
        para_format = node_pf.model_dump(exclude_none=True)
    else:
        para_format = node.meta.get("paragraph_format", {}) if node.meta else {}
    if not para_format:
        return
    pf = paragraph.paragraph_format
    if "alignment" in para_format and para_format["alignment"] is not None:
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            pf.alignment = WD_ALIGN_PARAGRAPH[para_format["alignment"].upper()]
        except (KeyError, TypeError):
            pass
    if "line_spacing" in para_format and para_format["line_spacing"] is not None:
        pf.line_spacing = float(para_format["line_spacing"])
    if "space_before" in para_format and para_format["space_before"] is not None:
        pf.space_before = Pt(float(para_format["space_before"]))
    if "space_after" in para_format and para_format["space_after"] is not None:
        pf.space_after = Pt(float(para_format["space_after"]))
    if "left_indent" in para_format and para_format["left_indent"] is not None:
        pf.left_indent = Pt(float(para_format["left_indent"]))
    if "right_indent" in para_format and para_format["right_indent"] is not None:
        pf.right_indent = Pt(float(para_format["right_indent"]))
    if "first_line_indent" in para_format and para_format["first_line_indent"] is not None:
        pf.first_line_indent = Pt(float(para_format["first_line_indent"]))
    if "page_break_before" in para_format and para_format["page_break_before"] is not None:
        pf.page_break_before = bool(para_format["page_break_before"])
    if "keep_lines_together" in para_format and para_format["keep_lines_together"] is not None:
        pf.keep_together = bool(para_format["keep_lines_together"])
    if "keep_with_next" in para_format and para_format["keep_with_next"] is not None:
        pf.keep_with_next = bool(para_format["keep_with_next"])


_LIST_NUM_CACHE: Dict[str, int] = {}
_LIST_NEXT_ID = 100


def _ensure_list_numbering(document: DocumentObject, is_ordered: bool, level: int) -> int:
    """Ensure a numbering definition exists for the given list type+level, return numId."""
    global _LIST_NEXT_ID
    cache_key = f"{'ordered' if is_ordered else 'bullet'}_{level}"
    if cache_key in _LIST_NUM_CACHE:
        return _LIST_NUM_CACHE[cache_key]

    numbering_part = document.part.numbering_part
    numbering_el = numbering_part._element

    abstract_num_id = _LIST_NEXT_ID
    num_id = _LIST_NEXT_ID
    _LIST_NEXT_ID += 1

    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Create abstractNum
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(level))

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if not is_ordered else "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if not is_ordered else "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(720 * (level + 1)))
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)
    lvl.append(pPr)

    abstract_num.append(lvl)
    numbering_el.append(abstract_num)

    # Create num
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id_ref = OxmlElement("w:abstractNumId")
    abstract_num_id_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_num_id_ref)
    numbering_el.append(num)

    _LIST_NUM_CACHE[cache_key] = num_id
    return num_id


def _apply_list_numbering(paragraph, num_id: int, ilvl: int) -> None:
    """Apply numPr to a paragraph using the given numbering definition."""
    pPr = paragraph._element.pPr
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._element.insert(0, pPr)

    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)

    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_el)

    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(num_id_el)


def _clear_story_body(story) -> None:
    """Remove existing content from a header or footer."""
    body = story._element
    for child in list(body):
        body.remove(child)


def _render_story_content(story, node: DocNode, image_blob_cache: Dict[str, bytes], style_map: Dict[str, str]) -> None:
    """Render a single paragraph or table node into a header/footer."""
    if node.type == NodeType.PARAGRAPH:
        if node.inline_content:
            para = _add_paragraph_with_inline_content_into_story(story, node, image_blob_cache)
        else:
            para = story.add_paragraph()
            if node.text:
                para.add_run(node.text)
        _apply_paragraph_formatting(para, node)
    elif node.type == NodeType.TABLE:
        meta = cast(Mapping[str, object], node.meta)
        _append_table_to_story(story, meta)


def _add_paragraph_with_inline_content_into_story(story, node: DocNode, image_blob_cache: Dict[str, bytes]) -> object:
    """Add a paragraph with inline content to a header/footer story."""
    paragraph = story.add_paragraph()
    for element in (node.inline_content or []):
        if isinstance(element, TextRun):
            run = paragraph.add_run(element.text)
            fmt = element.format.model_dump(exclude_none=True) if element.format else {}
            _apply_run_format(run, fmt)
        elif isinstance(element, HyperlinkRun):
            for child_run_data in element.children:
                run = paragraph.add_run(child_run_data.text)
                child_fmt = child_run_data.format.model_dump(exclude_none=True) if child_run_data.format else {}
                _apply_run_format(run, child_fmt)
        elif isinstance(element, InlineImageInline):
            blob = image_blob_cache.get(element.relationship_id)
            if blob is not None:
                paragraph.add_run().add_picture(BytesIO(blob), width=DEFAULT_IMAGE_RENDER_WIDTH)
    return paragraph


def _apply_section_properties(document: DocumentObject, sections_meta: list[dict]) -> None:
    """Apply section properties (page size, margins, orientation) to the template."""
    for s_idx, sec_meta in enumerate(sections_meta):
        if s_idx >= len(document.sections):
            break
        section = document.sections[s_idx]
        if "page_width" in sec_meta and sec_meta["page_width"] is not None:
            section.page_width = int(float(str(sec_meta["page_width"])))
        if "page_height" in sec_meta and sec_meta["page_height"] is not None:
            section.page_height = int(float(str(sec_meta["page_height"])))
        if "orientation" in sec_meta and sec_meta["orientation"] is not None:
            orient = str(sec_meta["orientation"])
            if "LANDSCAPE" in orient.upper():
                from docx.enum.section import WD_ORIENT
                section.orientation = WD_ORIENT.LANDSCAPE
        if "left_margin" in sec_meta and sec_meta["left_margin"] is not None:
            section.left_margin = int(float(str(sec_meta["left_margin"])))
        if "right_margin" in sec_meta and sec_meta["right_margin"] is not None:
            section.right_margin = int(float(str(sec_meta["right_margin"])))
        if "top_margin" in sec_meta and sec_meta["top_margin"] is not None:
            section.top_margin = int(float(str(sec_meta["top_margin"])))
        if "bottom_margin" in sec_meta and sec_meta["bottom_margin"] is not None:
            section.bottom_margin = int(float(str(sec_meta["bottom_margin"])))
        if sec_meta.get("different_first_page"):
            section.different_first_page_header_footer = True


def _append_table_to_story(story, meta: Mapping[str, object]) -> None:
    """Add a table to a header/footer story."""
    rows_raw = meta.get("rows", [])
    if not isinstance(rows_raw, Sequence) or isinstance(rows_raw, (str, bytes)) or not rows_raw:
        story.add_table(rows=1, cols=1)
        return
    ncols = max((len(r) for r in rows_raw if isinstance(r, Sequence) and not isinstance(r, (str, bytes))), default=1)
    nrows = len(rows_raw)
    table = story.add_table(rows=nrows, cols=ncols)
    for i, row in enumerate(rows_raw):
        if isinstance(row, Sequence) and not isinstance(row, (str, bytes)):
            for j, cell_text in enumerate(row):
                if j < ncols:
                    table.cell(i, j).text = str(cell_text) if cell_text else ""


def render_tree_to_template(
    root: DocNode,
    template_path: Union[str, Path],
    output_path: Union[str, Path],
    *,
    style_map: Optional[Dict[str, str]] = None,
    source_path: Optional[Union[str, Path]] = None,
) -> Path:
    """Render supported nodes into a template-driven DOCX document.

    When ``source_path`` matches the original DOCX that produced ``root``, IMAGE nodes whose
    ``meta`` exposes ``relationship_id`` can reuse embedded image blobs.
    """
    template = Path(template_path)
    output = Path(output_path)
    document = Document(template)
    resolved_style_map = {**DEFAULT_STYLE_MAP, **(style_map or {})}

    image_blob_cache: Dict[str, bytes] = {}
    if source_path is not None:
        src = Document(Path(source_path))
        image_blob_cache = extract_image_blob_cache(src)

    _clear_document_body(document)

    # Render headers and footers first
    for container in root.children:
        if container.type == NodeType.HEADER:
            section_idx = container.meta.get("section_index", 0)
            hdr_type = container.meta.get("header_type", "default")
            if section_idx >= len(document.sections):
                continue
            section = document.sections[section_idx]
            if hdr_type == "first":
                target = section.first_page_header
            else:
                target = section.header
            if target is not None:
                _clear_story_body(target)
                for child in container.children:
                    _render_story_content(target, child, image_blob_cache, resolved_style_map)
        elif container.type == NodeType.FOOTER:
            section_idx = container.meta.get("section_index", 0)
            hdr_type = container.meta.get("header_type", "default")
            if section_idx >= len(document.sections):
                continue
            section = document.sections[section_idx]
            if hdr_type == "first":
                target = section.first_page_footer
            else:
                target = section.footer
            if target is not None:
                _clear_story_body(target)
                for child in container.children:
                    _render_story_content(target, child, image_blob_cache, resolved_style_map)

    skipped_nodes: List[str] = []
    for node in walk_depth_first(root, skip_root=True):
        if node.type in (NodeType.HEADER, NodeType.FOOTER):
            continue
        if node.type == NodeType.HEADING:
            if node.inline_content:
                paragraph = _add_paragraph_with_inline_content(document, node, image_blob_cache)
            else:
                paragraph = document.add_paragraph(node.text or "")
            _safe_apply_style(paragraph, _heading_style_name(max(node.level, 1), resolved_style_map))
            _apply_paragraph_formatting(paragraph, node)
        elif node.type == NodeType.PARAGRAPH:
            if node.inline_content:
                paragraph = _add_paragraph_with_inline_content(document, node, image_blob_cache)
            else:
                paragraph = document.add_paragraph(node.text or "")
            _safe_apply_style(paragraph, resolved_style_map["paragraph"])
            _apply_paragraph_formatting(paragraph, node)
            if node.list_info is not None:
                num_id = _ensure_list_numbering(
                    document, node.list_info.is_ordered, node.list_info.level
                )
                _apply_list_numbering(paragraph, num_id, node.list_info.level)
        elif node.type == NodeType.TABLE:
            meta = cast(Mapping[str, object], node.meta)
            _append_table(document, meta)
        elif node.type == NodeType.IMAGE:
            rid_raw = node.meta.get("relationship_id")
            blob: Optional[bytes] = None
            if isinstance(rid_raw, str):
                blob = image_blob_cache.get(rid_raw)
            if blob is not None:
                _append_image(document, blob)
            else:
                skipped_nodes.append(
                    "image:no_blob_or_source"
                    if not source_path
                    else f"image:missing_rid={rid_raw}"
                )

    # Apply section properties
    sections_meta = root.meta.get("sections", [])
    if sections_meta:
        _apply_section_properties(document, sections_meta)

    if skipped_nodes:
        document.core_properties.comments = (
            "Skipped unsupported nodes during render: " + ", ".join(sorted(set(skipped_nodes)))
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def transform_docx(
    input_path: Union[str, Path],
    template_path: Union[str, Path],
    output_path: Union[str, Path],
    *,
    style_map: Optional[Dict[str, str]] = None,
) -> Path:
    """Parse a source document and render supported nodes into a template."""
    root = parse_docx(input_path)
    return render_tree_to_template(
        root,
        template_path,
        output_path,
        style_map=style_map,
        source_path=input_path,
    )
